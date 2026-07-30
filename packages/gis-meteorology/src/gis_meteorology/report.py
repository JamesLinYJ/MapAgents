# +-------------------------------------------------------------------------
#
#   地理智能平台 - 气象 DOCX 报告生成
#
#   文件:       report.py
#
#   日期:       2026年06月25日
#   作者:       JamesLinYJ
#   协助:       OpenAI Codex:GPT-5.5
# --------------------------------------------------------------------------

# 模块职责
#
# 将已经校验过的气象 metadata、统计事实和模型解读引用写入 DOCX。
# 这里不调用模型、不补造分析结论；缺少显式解读时由上游服务直接失败。

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any


def write_meteorological_report_docx(
    *,
    output_path: Path,
    dataset_id: str | None,
    filename: str,
    metadata: dict[str, Any],
    stats_rows: list[dict[str, Any]],
    llm_interpretation: str,
    generated_at: str,
) -> dict[str, Any]:
    """Write a compact DOCX report from explicit meteorological facts."""

    if not llm_interpretation.strip():
        raise ValueError("气象报告缺少经过校验的模型解读引用。")
    docx = _docx()
    document = docx.Document()
    document.add_heading("气象数据分析报告", level=0)
    document.add_paragraph(f"生成时间：{generated_at}")
    document.add_paragraph(f"数据文件：{filename}")
    if dataset_id:
        document.add_paragraph(f"数据集 ID：{dataset_id}")

    document.add_heading("一、数据概况", level=1)
    overview = document.add_table(rows=1, cols=2)
    overview.style = "Table Grid"
    overview.rows[0].cells[0].text = "项目"
    overview.rows[0].cells[1].text = "内容"
    _append_row(overview, "格式", str(metadata.get("format") or metadata.get("engine") or "未知"))
    _append_row(overview, "地理范围", _format_bounds(metadata.get("bounds")))
    _append_row(overview, "变量数量", str(len(_variables(metadata))))
    warnings = metadata.get("warnings")
    if isinstance(warnings, list) and warnings:
        _append_row(overview, "注意事项", "；".join(str(item) for item in warnings[:5]))

    document.add_heading("二、变量清单", level=1)
    variables = _variables(metadata)
    if variables:
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["变量", "单位", "维度", "值域", "地图能力"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for variable in variables:
            _append_variable_row(table, variable)
    else:
        document.add_paragraph("metadata 中未包含可展示变量。")

    document.add_heading("三、统计摘要", level=1)
    usable_stats = [row for row in stats_rows if isinstance(row, dict)]
    if usable_stats:
        table = document.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        headers = ["变量", "最小值", "最大值", "平均值", "P50", "P90", "说明"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for row in usable_stats:
            cells = table.add_row().cells
            cells[0].text = str(row.get("variable") or "")
            cells[1].text = _number(row.get("min"))
            cells[2].text = _number(row.get("max"))
            cells[3].text = _number(row.get("mean"))
            cells[4].text = _number(row.get("p50") or row.get("median"))
            cells[5].text = _number(row.get("p90"))
            cells[6].text = str(row.get("error") or row.get("unit") or "")
    else:
        document.add_paragraph("未生成变量统计摘要。")

    document.add_heading("四、模型解读", level=1)
    for paragraph in [item.strip() for item in llm_interpretation.splitlines() if item.strip()]:
        document.add_paragraph(paragraph)

    _prepare_document_for_save(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return {
        "outputPath": str(output_path),
        "filename": filename,
        "datasetId": dataset_id,
        "generatedAt": generated_at,
        "variableCount": len(variables),
        "statsRowCount": len(usable_stats),
    }


def write_nowcast_automation_report_docx(
    *,
    output_path: Path,
    automation_run_id: str,
    automation_id: str,
    automation_revision: int,
    started_at: str,
    completed_at: str,
    answer: str,
    analysis: dict[str, Any],
    artifacts: list[dict[str, Any]],
) -> dict[str, Any]:
    """Write a deterministic three-hour nowcast report from one persisted run."""

    regions = analysis.get("regions")
    if not isinstance(regions, list) or not regions:
        raise ValueError("短临自动化运行没有区域时间序列。")
    timelines = [
        (region, frame)
        for region in regions
        if isinstance(region, dict)
        for frame in region.get("timeline", [])
        if isinstance(frame, dict)
    ]
    if not timelines:
        raise ValueError("短临自动化运行没有可用于报告的时次统计。")
    peak_region, peak_frame = max(timelines, key=lambda item: _required_stat(item[1], "max"))
    peak_value = _required_stat(peak_frame, "max")
    generated_at = datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")

    docx = _docx()
    document = docx.Document()
    document.add_heading("三小时气象短临监测报告", level=0)
    document.add_paragraph(f"报告生成时间：{generated_at}")

    document.add_heading("一、运行依据", level=1)
    provenance = document.add_table(rows=1, cols=2)
    provenance.style = "Table Grid"
    provenance.rows[0].cells[0].text = "项目"
    provenance.rows[0].cells[1].text = "内容"
    _append_row(provenance, "自动化运行记录", automation_run_id)
    _append_row(provenance, "自动化定义", f"{automation_id}（修订 {automation_revision}）")
    _append_row(provenance, "执行时间", f"{started_at} 至 {completed_at}")
    _append_row(provenance, "分析变量", str(analysis.get("variable") or ""))
    _append_row(provenance, "序列标识", str(analysis.get("sequence_id") or ""))
    scope = analysis.get("scope")
    if isinstance(scope, dict):
        _append_row(provenance, "分析范围", str(scope.get("label") or scope.get("type") or ""))
    _append_row(provenance, "时次数量", str(len(timelines)))

    document.add_heading("二、自动化结论", level=1)
    document.add_paragraph(answer.strip())
    peak_time = _display_time(peak_frame.get("valid_time"))
    document.add_paragraph(
        f"全局最大 QPF 为 {peak_value:.3f} mm，出现在 {peak_time}"
        f"（相对 {int(peak_frame.get('lead_minutes') or 0)} 分钟，区域：{peak_region.get('label') or ''}）。"
    )

    document.add_heading("三、代表时次统计", level=1)
    for region in regions:
        if not isinstance(region, dict):
            continue
        timeline = [frame for frame in region.get("timeline", []) if isinstance(frame, dict)]
        if not timeline:
            continue
        document.add_heading(str(region.get("label") or "分析区域"), level=2)
        table = document.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        headers = ["有效时间", "相对分钟", "最大 QPF", "平均 QPF", "P90", "降水覆盖率", "降水等级"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for frame in _representative_frames(timeline):
            _append_nowcast_frame(table, frame)
        diagnosis = region.get("diagnosis")
        if isinstance(diagnosis, dict):
            document.add_paragraph(
                "诊断："
                f"{diagnosis.get('summary') or ''}；趋势 {diagnosis.get('trend') or ''}；"
                f"起雨 {diagnosis.get('onset_lead_minutes')} 分钟；"
                f"峰值 {diagnosis.get('peak_lead_minutes')} 分钟；"
                f"结束 {diagnosis.get('end_lead_minutes') if diagnosis.get('end_lead_minutes') is not None else '本时段未出现'}。"
            )

    document.add_heading("四、全部时次明细", level=1)
    full_table = document.add_table(rows=1, cols=9)
    full_table.style = "Table Grid"
    full_headers = ["区域", "序号", "有效时间", "相对分钟", "最大值", "平均值", "P90", "覆盖率", "文件"]
    for index, header in enumerate(full_headers):
        full_table.rows[0].cells[index].text = header
    for region, frame in timelines:
        stats = _frame_stats(frame)
        cells = full_table.add_row().cells
        cells[0].text = str(region.get("label") or "")
        cells[1].text = str(frame.get("sequence_index") or 0)
        cells[2].text = _display_time(frame.get("valid_time"))
        cells[3].text = str(frame.get("lead_minutes") or 0)
        cells[4].text = f"{float(stats.get('max') or 0):.3f}"
        cells[5].text = f"{float(stats.get('mean') or 0):.3f}"
        cells[6].text = f"{float(stats.get('p90') or 0):.3f}"
        cells[7].text = f"{float(stats.get('rain_coverage') or 0) * 100:.2f}%"
        cells[8].text = str(frame.get("filename") or "")

    document.add_heading("五、雨带移动与地图产物", level=1)
    movement = analysis.get("movement")
    if isinstance(movement, dict) and movement.get("available"):
        direction = str(movement.get("direction") or "未知方向").strip()
        movement_text = direction if direction.startswith("向") else f"向{direction}"
        document.add_paragraph(
            f"雨带{movement_text}移动约 "
            f"{float(movement.get('distance_km') or 0):.2f} km。"
        )
    else:
        document.add_paragraph("本次分析未生成可用的雨带移动诊断。")
    if artifacts:
        artifact_table = document.add_table(rows=1, cols=4)
        artifact_table.style = "Table Grid"
        for index, header in enumerate(["产物 ID", "名称", "类型", "访问地址"]):
            artifact_table.rows[0].cells[index].text = header
        for artifact in artifacts:
            cells = artifact_table.add_row().cells
            cells[0].text = str(artifact.get("artifact_id") or "")
            cells[1].text = str(artifact.get("name") or "")
            cells[2].text = str(artifact.get("artifact_type") or "")
            cells[3].text = str(artifact.get("uri") or "")
    else:
        document.add_paragraph("该运行没有关联地图或下载产物。")

    warnings = analysis.get("warnings")
    if isinstance(warnings, list) and warnings:
        document.add_heading("六、限制与注意事项", level=1)
        for warning in warnings:
            document.add_paragraph(str(warning), style="List Bullet")

    _prepare_document_for_save(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return {
        "outputPath": str(output_path),
        "automationRunId": automation_run_id,
        "generatedAt": generated_at,
        "frameCount": len(timelines),
        "regionCount": len(regions),
        "artifactCount": len(artifacts),
        "globalMaxQpf": peak_value,
        "globalMaxValidTime": str(peak_frame.get("valid_time") or ""),
        "globalMaxLeadMinutes": int(peak_frame.get("lead_minutes") or 0),
    }


def _frame_stats(frame: dict[str, Any]) -> dict[str, Any]:
    stats = frame.get("stats")
    if not isinstance(stats, dict):
        raise ValueError("短临时次缺少统计对象。")
    return stats


def _required_stat(frame: dict[str, Any], key: str) -> float:
    value = _frame_stats(frame).get(key)
    if not isinstance(value, (int, float)):
        raise ValueError(f"短临时次缺少数值统计：{key}。")
    return float(value)


def _representative_frames(timeline: list[dict[str, Any]]) -> list[dict[str, Any]]:
    selected = timeline[::6]
    if selected[-1] is not timeline[-1]:
        selected.append(timeline[-1])
    return selected


def _append_nowcast_frame(table: Any, frame: dict[str, Any]) -> None:
    stats = _frame_stats(frame)
    cells = table.add_row().cells
    cells[0].text = _display_time(frame.get("valid_time"))
    cells[1].text = str(frame.get("lead_minutes") or 0)
    cells[2].text = f"{float(stats.get('max') or 0):.3f} mm"
    cells[3].text = f"{float(stats.get('mean') or 0):.3f} mm"
    cells[4].text = f"{float(stats.get('p90') or 0):.3f} mm"
    cells[5].text = f"{float(stats.get('rain_coverage') or 0) * 100:.2f}%"
    cells[6].text = str(frame.get("rain_level") or "")


def _display_time(value: Any) -> str:
    text = str(value or "")
    if "T" in text:
        date, time = text.split("T", 1)
        return f"{date} {time[:5]}"
    return text


def _prepare_document_for_save(document: Any) -> None:
    """Normalize settings required by strict OOXML validators."""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings = document.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is None:
        zoom = OxmlElement("w:zoom")
        settings.insert(0, zoom)
    zoom.set(qn("w:val"), "bestFit")
    zoom.set(qn("w:percent"), "100")


def _append_row(table: Any, key: str, value: str) -> None:
    cells = table.add_row().cells
    cells[0].text = key
    cells[1].text = value


def _append_variable_row(table: Any, variable: dict[str, Any]) -> None:
    cells = table.add_row().cells
    cells[0].text = str(variable.get("name") or "")
    cells[1].text = str(variable.get("unit") or "")
    dimensions = variable.get("dimensions")
    cells[2].text = ", ".join(str(item) for item in dimensions) if isinstance(dimensions, list) else ""
    cells[3].text = _format_range(variable.get("valueRange"))
    cells[4].text = "可制图" if variable.get("mapReady") else "不可制图"


def _variables(metadata: dict[str, Any]) -> list[dict[str, Any]]:
    raw = metadata.get("variables")
    if not isinstance(raw, list):
        return []
    return [item for item in raw if isinstance(item, dict)]


def _format_bounds(value: Any) -> str:
    if not isinstance(value, list) or len(value) != 4:
        return "未提供"
    return ", ".join(_number(item) for item in value)


def _format_range(value: Any) -> str:
    if not isinstance(value, list) or len(value) != 2:
        return ""
    return f"{_number(value[0])} ~ {_number(value[1])}"


def _number(value: Any) -> str:
    try:
        return f"{float(value):.4g}"
    except (TypeError, ValueError):
        return ""


def _docx() -> Any:
    try:
        import docx

        return docx
    except Exception as exc:  # noqa: BLE001 - dependency import failure must surface clearly.
        raise RuntimeError("python-docx 不可用，无法生成气象 DOCX 报告。") from exc
